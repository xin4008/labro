import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.config import get_settings, resolve_path
from app.models.experiment import Experiment, ExperimentStep, FieldType

DEFAULT_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
PLACEHOLDER_ANALYSIS = "（请自行撰写）"
PLACEHOLDER_QUESTIONS = "（请自行撰写）"


def _set_run_font(
    run,
    font_name: str = DEFAULT_FONT,
    western: str = WESTERN_FONT,
    font_size=None,
    bold=None,
    color=None,
):
    run.font.name = western
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), western)
    r_fonts.set(qn("w:hAnsi"), western)
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text: str, bold: bool = False, size=Pt(12)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, font_size=size, bold=bold)
    return p


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, font_size=Pt(14 if level == 1 else 12), bold=True)
    return h


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge, val in kwargs.items():
        tag = f"w:{edge}"
        element = tc_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_pr.append(element)
        for key, value in val.items():
            element.set(qn(f"w:{key}"), str(value))


def _apply_three_line_table(table) -> None:
    """三线表：顶线、表头下线、底线加粗，无竖线。"""
    thick = {"val": "single", "sz": "12", "color": "000000", "space": "0"}
    thin = {"val": "single", "sz": "6", "color": "000000", "space": "0"}
    nil = {"val": "nil", "sz": "0", "color": "auto", "space": "0"}

    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            if ri == 0:
                _set_cell_border(
                    cell,
                    top=thick,
                    bottom=thin,
                    start=nil,
                    end=nil,
                )
            elif ri == len(table.rows) - 1:
                _set_cell_border(
                    cell,
                    top=nil,
                    bottom=thick,
                    start=nil,
                    end=nil,
                )
            else:
                _set_cell_border(
                    cell,
                    top=nil,
                    bottom=nil,
                    start=nil,
                    end=nil,
                )


def _set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    _set_run_font(run, font_size=Pt(10.5), bold=bold)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _resolve_image_path(experiment_id: str, image_path: str) -> Optional[Path]:
    if not image_path:
        return None
    parts = [p for p in image_path.split("/") if p]
    if "steps" in parts:
        si = parts.index("steps")
        if si + 2 < len(parts):
            step_id = parts[si + 1]
            name = parts[si + 2]
            path = (
                resolve_path(get_settings().storage.uploads_dir)
                / experiment_id
                / "steps"
                / step_id
                / name
            )
            if path.exists():
                return path
    return None


def _safe_filename(title: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", title).strip()
    return (cleaned or "实验报告")[:60]


def export_experiment_to_docx(experiment: Experiment) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(experiment.title)
    _set_run_font(tr, font_size=Pt(16), bold=True)

    date_str = ""
    if experiment.created_at:
        date_str = experiment.created_at.strftime("%Y年%m月%d日")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"实验日期：{date_str}")
    _set_run_font(mr, font_size=Pt(11))
    doc.add_paragraph()

    _heading(doc, "一、实验目的")
    _para(doc, experiment.purpose or "（未填写）")

    if experiment.reagents_instruments:
        _heading(doc, "试剂与仪器", level=2)
        for item in experiment.reagents_instruments:
            _para(doc, f"· {item}")

    _heading(doc, "二、实验步骤")
    steps = sorted(experiment.steps, key=lambda s: s.order_index)
    for idx, step in enumerate(steps, start=1):
        _heading(doc, f"步骤 {idx}：{step.title}", level=2)
        _para(doc, f"操作说明：{step.instructions or '—'}")

    _heading(doc, "三、实验现象")
    has_phenomenon = False
    for idx, step in enumerate(steps, start=1):
        texts: list[str] = []
        images: list[Path] = []
        for rec in step.records:
            if rec.field_type == FieldType.TEXT and rec.text_value:
                texts.append(f"{rec.field_label}：{rec.text_value}")
            if rec.field_type == FieldType.IMAGE and rec.image_path:
                p = _resolve_image_path(experiment.id, rec.image_path)
                if p and p.exists():
                    images.append(p)
        if step.expected_phenomenon:
            texts.insert(0, f"［文献预期］{step.expected_phenomenon}")
        if texts or images:
            has_phenomenon = True
            _para(doc, f"步骤 {idx}", bold=True)
            for t in texts:
                _para(doc, t)
            for img in images:
                try:
                    doc.add_picture(str(img), width=Inches(4.5))
                except Exception:
                    _para(doc, f"[图片：{img.name}]")

    if not has_phenomenon:
        _para(doc, "（实验过程中未记录现象描述或图片）")

    _heading(doc, "四、实验数据")
    has_data = False
    for idx, step in enumerate(steps, start=1):
        table_records = [r for r in step.records if r.field_type == FieldType.NUMBER]
        if not table_records:
            continue
        has_data = True
        _para(doc, f"步骤 {idx}：{step.title}", bold=True)

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], "数据项", bold=True)
        _set_cell_text(hdr[1], "数值/内容", bold=True)
        _set_cell_text(hdr[2], "单位", bold=True)

        for rec in table_records:
            if rec.field_type == FieldType.IMAGE:
                continue
            row = table.add_row().cells
            _set_cell_text(row[0], rec.field_label)
            if rec.field_type == FieldType.NUMBER:
                val = "" if rec.number_value is None else str(rec.number_value)
                _set_cell_text(row[1], val)
            else:
                _set_cell_text(row[1], rec.text_value or "")
            _set_cell_text(row[2], rec.unit or "—")

        _apply_three_line_table(table)
        doc.add_paragraph()

    if not has_data:
        _para(doc, "（未记录数值型实验数据）")

    _heading(doc, "五、分析与讨论")
    p = _para(doc, PLACEHOLDER_ANALYSIS)
    for run in p.runs:
        run.italic = True

    _heading(doc, "六、思考题")
    p = _para(doc, PLACEHOLDER_QUESTIONS)
    for run in p.runs:
        run.italic = True

    settings = get_settings()
    export_dir = resolve_path(settings.storage.exports_dir)
    export_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{_safe_filename(experiment.title)}_{stamp}.docx"
    filepath = export_dir / filename
    doc.save(str(filepath))
    return filepath
