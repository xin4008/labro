import re
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import fitz
import httpx
from docx import Document

# 需要登录、禁止爬虫的站点 → 友好提示
_BLOCKED_DOMAIN_HINTS: dict[str, str] = {
    "cnki.net": "中国知网需要机构账号登录，程序无法直接打开知网链接。请在学校网络下从知网下载 PDF，再用「上传 PDF」导入。",
    "wanfangdata.com": "万方数据需要登录访问，请下载 PDF 后上传。",
    "cqvip.com": "维普资讯需要登录访问，请下载 PDF 后上传。",
    "sciencedirect.com": "ScienceDirect 需要订阅权限，请下载 PDF 后上传。",
    "springer.com": "Springer 需要订阅权限，请下载 PDF 后上传。",
    "wiley.com": "Wiley 需要订阅权限，请下载 PDF 后上传。",
}

_DEFAULT_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip = False
        if tag in {"p", "div", "br", "li", "h1", "h2", "h3", "tr"}:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip and data.strip():
            self._chunks.append(data.strip())

    def get_text(self) -> str:
        text = " ".join(self._chunks)
        return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_text_from_pdf(path: Path) -> str:
    doc = fitz.open(path)
    parts: list[str] = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts).strip()


def extract_text_from_docx(path: Path) -> str:
    document = Document(path)
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text_from_bytes(content: bytes, doc_type: str) -> str:
    if doc_type == "pdf":
        doc = fitz.open(stream=content, filetype="pdf")
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(parts).strip()
    if doc_type == "docx":
        document = Document(BytesIO(content))
        return "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
    raise ValueError(f"不支持的文档类型：{doc_type}")


def _friendly_fetch_error(url: str, status_code: int) -> str:
    host = urlparse(url).netloc.lower()
    for domain, hint in _BLOCKED_DOMAIN_HINTS.items():
        if domain in host:
            return hint
    if status_code == 403:
        return (
            "该网站拒绝自动访问（403 禁止访问），常见于知网、数据库等需登录页面。"
            "请改用手动下载 PDF/DOCX 后上传。"
        )
    if status_code == 401:
        return "该页面需要登录才能查看，请下载 PDF 后上传。"
    if status_code == 404:
        return "链接不存在或已失效（404），请检查网址是否正确。"
    return f"无法访问该网址（HTTP {status_code}），请下载 PDF 后上传。"


async def fetch_url_text(url: str, timeout: float = 30.0) -> tuple[str, str]:
    """返回 (提取的文本, 实际文件名提示)。"""
    headers = {"User-Agent": _DEFAULT_USER_AGENT, "Accept-Language": "zh-CN,zh;q=0.9"}
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
        try:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            code = exc.response.status_code
            raise ValueError(_friendly_fetch_error(url, code)) from exc
        except httpx.RequestError as exc:
            raise ValueError(f"网络连接失败：{exc}。请检查网络或改用手动上传 PDF。") from exc

        content_type = response.headers.get("content-type", "").lower()

        if "pdf" in content_type or url.lower().endswith(".pdf"):
            text = extract_text_from_bytes(response.content, "pdf")
            return text, "webpage.pdf"

        if (
            "wordprocessingml" in content_type
            or url.lower().endswith(".docx")
        ):
            text = extract_text_from_bytes(response.content, "docx")
            return text, "webpage.docx"

        html = response.text
        parser = _HTMLTextExtractor()
        parser.feed(html)
        return parser.get_text(), "webpage.html"


def truncate_text(text: str, max_chars: int = 80000) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[文本已截断，仅发送前部分内容供 AI 解析]"
