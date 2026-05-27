import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
DEFAULT_FONT = "宋体"


def _set_run_font(run, font_name=DEFAULT_FONT, font_size=None, bold=None):
    """Apply a CJK-capable font to a text run."""
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.bold = bold


def _apply_font_to_paragraph(paragraph, font_name=DEFAULT_FONT, font_size=None):
    for run in paragraph.runs:
        _set_run_font(run, font_name, font_size=font_size)


def _set_cell_text(cell, text, font_name=DEFAULT_FONT):
    cell.text = text
    _apply_font_to_paragraph(cell.paragraphs[0], font_name)


def _configure_document_defaults(doc, font_name=DEFAULT_FONT):
    """Set document-wide Chinese language and font defaults."""
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        return

    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "zh-CN")


def _configure_style_fonts(doc, font_name=DEFAULT_FONT):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = font_name
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)


def _apply_fonts_to_document(doc, font_name=DEFAULT_FONT):
    for paragraph in doc.paragraphs:
        _apply_font_to_paragraph(paragraph, font_name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_font_to_paragraph(paragraph, font_name)


def export_to_docx(experiment):
    """Generate a .docx file for the given experiment and return the file path."""
    doc = Document()
    _configure_document_defaults(doc)
    _configure_style_fonts(doc)

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_heading(experiment["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"日期：{experiment['created_at']}")
    _set_run_font(meta_run, font_size=Pt(11))

    doc.add_paragraph()

    doc.add_heading("一、实验目的", level=1)
    doc.add_paragraph(experiment["objective"])

    doc.add_heading("二、实验步骤与记录", level=1)

    for step in experiment["steps"]:
        doc.add_heading(f"步骤 {step['id']}：{step['title']}", level=2)
        doc.add_paragraph(f"操作说明：{step['instruction']}")

        if step.get("data_fields"):
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            _set_cell_text(hdr[0], "数据项")
            _set_cell_text(hdr[1], "数值")
            _set_cell_text(hdr[2], "单位")
            for field in step["data_fields"]:
                row = table.add_row()
                _set_cell_text(row.cells[0], field["label"])
                _set_cell_text(row.cells[1], field["value"])
                _set_cell_text(row.cells[2], field["unit"])
            doc.add_paragraph()

        if step.get("observation"):
            doc.add_paragraph(f"实验现象：{step['observation']}")

        notes = step.get("ai_notes", {})
        if notes.get("precautions"):
            doc.add_paragraph(f"注意事项：{notes['precautions']}")
        if notes.get("prediction"):
            doc.add_paragraph(f"预测现象：{notes['prediction']}")
        if notes.get("safety"):
            p = doc.add_paragraph()
            run = p.add_run(f"安全提醒：{notes['safety']}")
            _set_run_font(run, bold=True)

        if step.get("photos"):
            for photo_rel in step["photos"]:
                exp_dir = _find_exp_dir(experiment["id"])
                photo_path = os.path.join(exp_dir, photo_rel)
                if os.path.isfile(photo_path):
                    try:
                        doc.add_picture(photo_path, width=Inches(4))
                    except Exception:
                        doc.add_paragraph(f"[照片：{photo_rel}]")

    doc.add_heading("三、参考文献", level=1)
    doc.add_paragraph(experiment["references"])

    _apply_fonts_to_document(doc)

    export_dir = os.path.join(DATA_DIR, "exports")
    os.makedirs(export_dir, exist_ok=True)
    safe_title = "".join(c for c in experiment["title"] if c.isalnum() or c in "._- ")[:50]
    filename = f"{safe_title}_{experiment['created_at']}.docx"
    filepath = os.path.join(export_dir, filename)
    doc.save(filepath)
    return filepath


def _find_exp_dir(exp_id):
    for d in os.listdir(DATA_DIR):
        if d.endswith(f"_{exp_id}"):
            full = os.path.join(DATA_DIR, d)
            if os.path.isdir(full):
                return full
    return None
